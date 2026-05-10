#!/usr/bin/env python3
"""Extract one Act.Chapter section from a .docx manuscript for ElevenLabs script generation."""

from __future__ import annotations

import argparse
import html
import json
import math
import os
import random
import re
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from openai import OpenAI

NARRATOR_NAME = "Kertoja"
LANGUAGE_ALIASES = {
    "fi": "finnish",
    "finnish": "finnish",
    "suomi": "finnish",
    "en": "english",
    "eng": "english",
    "english": "english",
    "es": "spanish",
    "spa": "spanish",
    "spanish": "spanish",
    "espanol": "spanish",
}
MAX_FRAGMENT_LEN = 4000
SINGLE_VOICE_MAX_FRAGMENT_LEN = 4095
INTER_CHUNK_BREAK = "0.4s"
INTER_LINE_BREAK = "0.8s"
EMOJI_RE = re.compile("[\U0001F300-\U0001FAFF\U00002700-\U000027BF\U000024C2-\U0001F251]+", flags=re.UNICODE)
CHAT_LINE_RE = re.compile(r"^\s*\[([^\]\n]{1,64})\]\s*:\s*(.+?)\s*$")

SPEECH_VERB_ROOTS = {
    "finnish": ["sano", "kysy", "vasta", "huuda", "totea", "juttele", "mainitse", "kommentoi", "kuiska", "karju", "mumise", "lausu", "selitä", "ilmoita", "myönnä", "kiistä", "toista", "mutise", "hihku", "naurahda", "neuvo", "täydentä", "jatka", "lisä", "toka", "huikka", "kerro"],
    "english": ["say", "ask", "answer", "reply", "shout", "whisper", "murmur", "remark", "state", "announce", "admit", "deny", "yell", "scream", "observe", "note", "repeat", "explain", "utter", "tell"],
    "spanish": ["dec", "pregunt", "respond", "contest", "grit", "susurr", "murmur", "afirm", "neg", "coment", "explic", "repet", "pronunci", "admit", "anunci", "observ", "añad", "indic", "declar", "dij"],
}
SPEECH_VERB_SUFFIXES = {
    "finnish": ["a", "an", "aa", "aan", "oi", "oin", "oivat", "ot", "omme", "ossa", "osta", "osti", "oivat", "oisi", "oisi", "oimme", "oivat", "n", "t", "mme", "tte", "vat", "in", "it", "imme", "itte", "ivat", "isin", "isit", "isi", "isimme", "isitte", "isivat", "nut", "nyt", "neet", "neet", "massa", "masta", "malla", "malle", "melta", "si"],
    "english": ["", "s", "ed", "ing", "er", "ers", "ly", "able", "ably", "ation", "ations", "ment", "ments", "t", "ts", "d", "en", "es"],
    "spanish": ["ar", "o", "as", "a", "amos", "áis", "an", "é", "aste", "ó", "aron", "aba", "abas", "ábamos", "aban", "aré", "arás", "ará", "arán", "aría", "arías", "arían", "ado", "ada", "ados", "adas", "ando", "en", "emos", "éis"],
}



def parse_bool_arg(value: str) -> bool:
    v = (value or "").strip().lower()
    if v in {"true", "yes", "1", "on"}:
        return True
    if v in {"false", "no", "0", "off"}:
        return False
    raise argparse.ArgumentTypeError("Boolean value expected: true/false/yes/no/1/0/on/off")
def _build_speech_verbs() -> set[str]:
    forms: set[str] = set()
    for lang, roots in SPEECH_VERB_ROOTS.items():
        for root in roots:
            for suffix in SPEECH_VERB_SUFFIXES[lang]:
                token = (root + suffix).lower().strip()
                if 3 <= len(token) <= 24:
                    forms.add(token)
    return forms

SPEECH_VERBS = _build_speech_verbs() | {"neuvonut", "täydentänyt", "huikkasi", "kysyi", "vastasi", "sanoi"}

@dataclass
class OpenAIProfile:
    profile_id: str
    voice: str
    model: str
    instructions: str


@dataclass
class NarratorConfig:
    voices: dict[str, str]
    aliases: dict[str, str]
    openai_profiles: dict[str, dict[str, OpenAIProfile]]


@dataclass
class ScriptSegment:
    segment_id: int
    type: str  # "narration" | "dialogue"
    speaker: str
    text: str
    confidence: float = 1.0


@dataclass
class ChatLine:
    nickname: str
    message: str
    raw: str


def clean_chat_nickname(nickname: str) -> str:
    nick = nickname.strip().strip("[]").strip("_")
    return nick or "anonymous"


def parse_chat_line(line: str) -> ChatLine | None:
    m = CHAT_LINE_RE.match(line or "")
    if not m:
        return None
    message = (m.group(2) or "").strip()
    if not message:
        return None
    return ChatLine(nickname=clean_chat_nickname(m.group(1)), message=message, raw=line)


def normalise_alias_key(value: str) -> str:
    value = value.strip().lower().replace("0", "o")
    return re.sub(r"[^a-zåäöáéíóúüñ0-9]+", "", value, flags=re.IGNORECASE)


def style_level(style_name: str) -> int | None:
    name = style_name.strip().lower()
    m = re.search(r"(\d+)$", name)
    if m:
        return int(m.group(1))
    if "otsikko 1" in name or "heading 1" in name:
        return 1
    if "otsikko 2" in name or "heading 2" in name:
        return 2
    return None


def load_narrator_config(path: Path) -> NarratorConfig:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        raise ValueError(f"Narrators-tiedosto ei ole validia JSONia: {e}") from e
    if not isinstance(data, dict):
        raise ValueError("Narrators-tiedoston pitää olla JSON-objekti")
    voices_raw = data.get("voices")
    if not isinstance(voices_raw, dict):
        raise ValueError("Narrators-tiedoston 'voices' pitää olla objekti")

    voices: dict[str, str] = {}
    openai_profiles: dict[str, dict[str, OpenAIProfile]] = {}
    for name, meta in voices_raw.items():
        if not isinstance(meta, dict):
            raise ValueError(f"Voice-entry '{name}' on virheellinen: pitää olla objekti")
        ids = meta.get("ids")
        if not isinstance(ids, dict):
            raise ValueError(f"Voice-entry '{name}' on virheellinen: 'ids' puuttuu tai ei ole objekti")
        eleven = ids.get("elevenlabs")
        if not eleven:
            raise ValueError(f"Voice-entry '{name}' on virheellinen: ids.elevenlabs puuttuu")
        narrator_name = str(name)
        voices[narrator_name] = str(eleven)
        raw_profiles = meta.get("openai_profiles")
        if raw_profiles is None:
            continue
        if not isinstance(raw_profiles, dict):
            raise ValueError(f"Voice-entry '{name}' on virheellinen: 'openai_profiles' pitää olla objekti")
        speaker_profiles: dict[str, OpenAIProfile] = {}
        for language_key, profile_meta in raw_profiles.items():
            if not isinstance(profile_meta, dict):
                raise ValueError(f"Voice-entry '{name}' profile '{language_key}' on virheellinen: pitää olla objekti")
            voice = str(profile_meta.get("voice", "")).strip()
            if not voice:
                raise ValueError(f"Voice-entry '{name}' profile '{language_key}' on virheellinen: 'voice' puuttuu")
            profile_id = str(profile_meta.get("profile_id", "")).strip() or f"{normalise_alias_key(narrator_name)}-{voice}-{language_key}"
            instructions_raw = profile_meta.get("instructions")
            if instructions_raw is None:
                print(f"Warning: OpenAI instructions missing for {narrator_name}/{language_key}; using empty instructions.", file=sys.stderr)
                instructions = ""
            else:
                instructions = str(instructions_raw)
            model = str(profile_meta.get("model", "gpt-4o-mini-tts")).strip() or "gpt-4o-mini-tts"
            speaker_profiles[str(language_key)] = OpenAIProfile(
                profile_id=profile_id,
                voice=voice,
                model=model,
                instructions=instructions,
            )
        if speaker_profiles:
            openai_profiles[narrator_name] = speaker_profiles

    if NARRATOR_NAME not in voices:
        raise ValueError(f"Narrators-tiedostossa pitää olla '{NARRATOR_NAME}'")

    alias_map = build_alias_map(voices)
    aliases = data.get("aliases", {})
    if aliases is not None:
        if not isinstance(aliases, dict):
            raise ValueError("Narrators-tiedoston 'aliases' pitää olla objekti")
        for canonical, values in aliases.items():
            canonical_name = str(canonical)
            if canonical_name not in voices:
                print(f"Warning: alias canonical voice missing, skipped: {canonical_name}", file=sys.stderr)
                continue
            if not isinstance(values, list):
                print(f"Warning: alias list invalid, skipped: {canonical_name}", file=sys.stderr)
                continue
            for alias in values:
                if isinstance(alias, str) and alias.strip():
                    alias_map[normalise_alias_key(alias)] = canonical_name
    return NarratorConfig(voices=voices, aliases=alias_map, openai_profiles=openai_profiles)


def normalise_language_key(language_name: str) -> str:
    key = (language_name or "").strip().lower()
    aliases = {
        "finnish": "fi",
        "suomi": "fi",
        "fi": "fi",
        "english": "en",
        "en": "en",
        "spanish": "es",
        "espanol": "es",
        "español": "es",
        "es": "es",
    }
    return aliases.get(key, key)


def resolve_openai_profile(
    speaker: str,
    language_name: str,
    cfg: NarratorConfig,
    variant: str | None = None,
) -> OpenAIProfile | None:
    speaker_profiles = cfg.openai_profiles.get(speaker, {})
    lang_key = normalise_language_key(language_name)
    variant_key = (variant or "").strip().lower()
    candidate_keys: list[str] = []
    if variant_key:
        candidate_keys.append(f"{lang_key}_{variant_key}")
    candidate_keys.append(lang_key)
    legacy_key = {"fi": "finnish", "en": "english", "es": "spanish"}.get(lang_key)
    if legacy_key:
        candidate_keys.append(legacy_key)
    for key in candidate_keys:
        if key in speaker_profiles:
            return speaker_profiles[key]
    if speaker_profiles:
        return next(iter(speaker_profiles.values()))
    narrator_profiles = cfg.openai_profiles.get(NARRATOR_NAME, {})
    for key in candidate_keys:
        if key in narrator_profiles:
            return narrator_profiles[key]
    if narrator_profiles:
        return next(iter(narrator_profiles.values()))
    return None


def build_alias_map(narrators: dict[str, str]) -> dict[str, str]:
    alias_map: dict[str, str] = {}
    ambiguous: set[str] = set()
    for canonical_name in narrators:
        if canonical_name.startswith("chat"):
            continue
        canonical_key = normalise_alias_key(canonical_name)
        alias_map[canonical_key] = canonical_name
        parts = [p for p in re.split(r"[\s_\-]+", canonical_name.strip()) if p]
        for part in parts:
            part_key = normalise_alias_key(part)
            if len(part_key) >= 3:
                if part_key not in alias_map:
                    alias_map[part_key] = canonical_name
                elif alias_map[part_key] != canonical_name:
                    ambiguous.add(part_key)
        if len(parts) >= 2:
            last_key = normalise_alias_key(parts[-1])
            if last_key not in alias_map:
                alias_map[last_key] = canonical_name
            elif alias_map[last_key] != canonical_name:
                ambiguous.add(last_key)
    for key in ambiguous:
        alias_map.pop(key, None)
    alias_map[normalise_alias_key(NARRATOR_NAME)] = NARRATOR_NAME
    return alias_map


def resolve_voice_name(candidate: str, narrators: dict[str, str], alias_map: dict[str, str]) -> str:
    raw = (candidate or "").strip()
    if not raw:
        return NARRATOR_NAME
    if raw in narrators:
        return raw
    if len(raw) > 48 or re.search(r"[,.;!?]", raw):
        return NARRATOR_NAME
    if len(re.findall(r"\w+", raw)) > 5:
        return NARRATOR_NAME
    return alias_map.get(normalise_alias_key(raw), NARRATOR_NAME)


def clean_dialogue_text(text: str) -> str:
    cleaned = (text or "").strip()
    # Remove common leading dialogue markers but keep dashes inside the utterance.
    cleaned = re.sub(r'^\s*[–—-]\s*', "", cleaned)
    cleaned = re.sub(r'^\s*[“"]\s*', "", cleaned)
    return cleaned.strip()


def ensure_sentence_punctuation(text: str) -> str:
    value = (text or "").strip()
    if not value:
        return value
    if value[-1] in ".!?…":
        return value
    return f"{value}."


def _dialogue_tag_subject_present(text: str, narrators: dict[str, str], alias_map: dict[str, str]) -> bool:
    lowered = f" {text.lower()} "
    if re.search(r"\bhän\b", lowered, re.IGNORECASE):
        return True
    for name in narrators:
        if name == NARRATOR_NAME:
            continue
        if re.search(rf"\b{re.escape(name.lower())}\b", lowered):
            return True
    for canonical in set(alias_map.values()):
        if canonical == NARRATOR_NAME:
            continue
        lowered_name = canonical.lower()
        if re.search(rf"\b{re.escape(lowered_name)}\b", lowered):
            return True
    return False


def looks_like_dialogue_tag(text: str, narrators: dict[str, str], alias_map: dict[str, str]) -> bool:
    candidate = (text or "").strip(" –—-\t")
    if not candidate:
        return False
    if not _dialogue_tag_subject_present(candidate, narrators, alias_map):
        return False
    words = re.findall(r"[A-Za-zÅÄÖåäöÁÉÍÓÚÜÑáéíóúüñ'-]+", candidate.lower())
    return any(word in SPEECH_VERBS for word in words)


def split_trailing_dialogue_tag(text: str, speaker: str, narrators: dict[str, str], alias_map: dict[str, str]) -> list[ScriptSegment]:
    content = clean_dialogue_text(text)
    clauses = [c.strip() for c in re.split(r"(?<=[.!?])\s+", content) if c.strip()]
    spoken_part = ""
    trailing = ""
    if len(clauses) >= 2:
        spoken_part = " ".join(clauses[:-1]).strip()
        trailing = clauses[-1].strip()
    else:
        comma_idx = content.rfind(",")
        if comma_idx <= 0:
            return [ScriptSegment(segment_id=0, type="dialogue", speaker=speaker, text=content)]
        spoken_part = content[:comma_idx + 1].strip()
        trailing = content[comma_idx + 1:].strip()
    if not spoken_part or not trailing or not looks_like_dialogue_tag(trailing, narrators, alias_map):
        return [ScriptSegment(segment_id=0, type="dialogue", speaker=speaker, text=content)]
    return [
        ScriptSegment(segment_id=0, type="dialogue", speaker=speaker, text=spoken_part),
        ScriptSegment(segment_id=0, type="narration", speaker=NARRATOR_NAME, text=ensure_sentence_punctuation(trailing)),
    ]


def split_embedded_dialogue_tag(text: str, speaker: str, narrators: dict[str, str], alias_map: dict[str, str]) -> list[ScriptSegment]:
    content = clean_dialogue_text(text)
    marker = re.search(r"\s[–—-]\s+", content)
    if not marker:
        return [ScriptSegment(segment_id=0, type="dialogue", speaker=speaker, text=content)]
    before = content[:marker.start()].strip()
    after = clean_dialogue_text(content[marker.end():])
    clauses = [c.strip() for c in re.split(r"(?<=[.!?])\s+", before) if c.strip()]
    spoken_head = ""
    embedded_tag = ""
    if len(clauses) >= 2:
        spoken_head = " ".join(clauses[:-1]).strip()
        embedded_tag = clauses[-1].strip()
    else:
        comma_idx = before.rfind(",")
        if comma_idx <= 0:
            return [ScriptSegment(segment_id=0, type="dialogue", speaker=speaker, text=content)]
        spoken_head = before[:comma_idx + 1].strip()
        embedded_tag = before[comma_idx + 1:].strip()
    if not spoken_head or not embedded_tag:
        return [ScriptSegment(segment_id=0, type="dialogue", speaker=speaker, text=content)]
    if not looks_like_dialogue_tag(embedded_tag, narrators, alias_map):
        return [ScriptSegment(segment_id=0, type="dialogue", speaker=speaker, text=content)]
    result = [
        ScriptSegment(segment_id=0, type="dialogue", speaker=speaker, text=spoken_head),
        ScriptSegment(segment_id=0, type="narration", speaker=NARRATOR_NAME, text=ensure_sentence_punctuation(embedded_tag)),
    ]
    if after:
        result.append(ScriptSegment(segment_id=0, type="dialogue", speaker=speaker, text=ensure_sentence_punctuation(after)))
    return result


def split_dialogue_with_embedded_tags(segment: ScriptSegment, narrators: dict[str, str], alias_map: dict[str, str]) -> list[ScriptSegment]:
    if segment.type != "dialogue":
        return [segment]
    embedded = split_embedded_dialogue_tag(segment.text, segment.speaker, narrators, alias_map)
    if len(embedded) > 1:
        return embedded
    trailing = split_trailing_dialogue_tag(segment.text, segment.speaker, narrators, alias_map)
    if len(trailing) > 1:
        return trailing
    return [ScriptSegment(segment_id=segment.segment_id, type="dialogue", speaker=segment.speaker, text=clean_dialogue_text(segment.text), confidence=segment.confidence)]


def normalise_segment_text_for_compare(text: str) -> str:
    value = clean_dialogue_text(text).lower().strip()
    value = re.sub(r'[“”"\'«»]', "", value)
    value = re.sub(r"\s+", " ", value).strip()
    value = value[:-1].rstrip() if value.endswith(".") else value
    return value


def deduplicate_segments(segments: list[ScriptSegment]) -> list[ScriptSegment]:
    if not segments:
        return []
    deduped: list[ScriptSegment] = [segments[0]]
    for seg in segments[1:]:
        prev = deduped[-1]
        if seg.type != "dialogue" or prev.type != "dialogue" or seg.speaker != prev.speaker:
            deduped.append(seg)
            continue
        prev_norm = normalise_segment_text_for_compare(prev.text)
        curr_norm = normalise_segment_text_for_compare(seg.text)
        if not prev_norm or not curr_norm:
            deduped.append(seg)
            continue
        same = prev_norm == curr_norm
        near_dup = prev_norm in curr_norm or curr_norm in prev_norm
        if not (same or near_dup):
            deduped.append(seg)
            continue
        prev_clean = clean_dialogue_text(prev.text)
        curr_clean = clean_dialogue_text(seg.text)
        if prev_norm in curr_norm and len(curr_clean) > len(prev_clean):
            deduped[-1] = seg
        elif curr_norm in prev_norm and len(prev_clean) >= len(curr_clean):
            continue
        elif curr_clean != prev_clean:
            deduped[-1] = seg
    return deduped


def normalise_segments(
    segments: list[ScriptSegment],
    narrators: dict[str, str],
    alias_map: dict[str, str],
) -> list[ScriptSegment]:
    normalised: list[ScriptSegment] = []
    for segment in segments:
        seg_type = "dialogue" if segment.type == "dialogue" else "narration"
        if seg_type == "narration":
            speaker = NARRATOR_NAME
            text = segment.text.strip()
        else:
            raw_speaker = segment.speaker
            speaker = resolve_voice_name(raw_speaker, narrators, alias_map)
            if speaker == NARRATOR_NAME and raw_speaker.strip() and normalise_alias_key(raw_speaker) != normalise_alias_key(NARRATOR_NAME):
                print(f'WARNING: Speaker "{raw_speaker}" was detected but no matching narrator was found. Falling back to {NARRATOR_NAME}.', file=sys.stderr)
            text = clean_dialogue_text(segment.text)
        if not text:
            continue
        cleaned = ScriptSegment(segment_id=segment.segment_id, type=seg_type, speaker=speaker, text=text, confidence=segment.confidence)
        if cleaned.type == "dialogue":
            split_segments = split_dialogue_with_embedded_tags(cleaned, narrators, alias_map)
            for idx, split_segment in enumerate(split_segments):
                normalised.append(
                    ScriptSegment(
                        segment_id=cleaned.segment_id * 10 + idx,
                        type=split_segment.type,
                        speaker=split_segment.speaker,
                        text=split_segment.text,
                        confidence=cleaned.confidence,
                    )
                )
        else:
            normalised.append(cleaned)
    return deduplicate_segments(normalised)

# rest mostly unchanged, shortened for brevity in this patch context
# (kept deterministic parser and chat handling)

def extract_section(docx_path: Path, content_id: str | None) -> tuple[str, int, int, str]:
    document = Document(str(docx_path))
    if not content_id:
        lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        return "\n\n".join(lines).strip() + "\n", 1, 1, docx_path.stem
    if not re.fullmatch(r"\d+\.\d+", content_id):
        raise ValueError("--content pitää olla muodossa ACT.LUKU, esim. 1.3")
    target_act, target_chapter = map(int, content_id.split("."))
    act = chapter = 0
    collecting = False
    lines: list[str] = []
    chapter_title = "luku"
    for para in document.paragraphs:
        text = para.text.strip()
        level = style_level(getattr(para.style, "name", ""))
        if level == 1:
            act += 1; chapter = 0
            if collecting: break
            continue
        if level == 2:
            chapter += 1
            if collecting: break
            collecting = act == target_act and chapter == target_chapter
            if collecting: chapter_title = text or f"Luku {chapter}"
            continue
        if collecting and text:
            lines.append(text)
    if not lines:
        raise LookupError(f"Sisältöä ei löytynyt kohdalle {content_id} tiedostosta {docx_path}")
    return "\n\n".join(lines).strip() + "\n", target_act, target_chapter, chapter_title

# Keep existing helper functions from original file.

def detect_language_name(text: str) -> str:
    lowered = text.lower(); tokens = re.findall(r"[a-zåäö]+", lowered)
    if not tokens: return "finnish"
    en = {"the", "and", "with", "viewers", "counter", "strike", "reaction", "chat"}
    es = {"la", "el", "de", "que", "y", "banda"}
    fi = {"on", "ja", "että", "oli", "mutta", "tunnettiin", "nimimerkillä"}
    en_score = sum(1 for t in tokens if t in en); es_score = sum(1 for t in tokens if t in es); fi_score = sum(1 for t in tokens if t in fi)
    fi_score += sum(1 for ch in lowered if ch in "åäö") * 0.6
    if math.isclose(en_score, es_score) and en_score > fi_score and " la " in f" {lowered} ": es_score += 0.2
    if fi_score >= en_score and fi_score >= es_score: return "finnish"
    return "english" if en_score >= es_score else "spanish"


def strip_non_prose_for_language_detection(text: str) -> str:
    cleaned = re.sub(r"<[^>]+>", " ", text or "")
    cleaned = re.sub(r"\b[a-z]+://\S+\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\b[\w.\-]+@[\w.\-]+\b", " ", cleaned)
    cleaned = re.sub(r"\b(?:[A-Za-z]:)?[/\\][^\s]+\b", " ", cleaned)
    cleaned = re.sub(r"\b[\w.\-]+\.(?:xml|json|txt|docx|md|wav|mp3|m4a)\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"^\s*[^:\n]{1,40}:\s*", " ", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def dominant_language_from_segments(segments: list[ScriptSegment]) -> str:
    weighted_scores = {"finnish": 0.0, "english": 0.0, "spanish": 0.0}
    for seg in segments:
        prose = strip_non_prose_for_language_detection(seg.text)
        if not prose:
            continue
        lang = detect_language_name(prose)
        alpha_chars = len(re.findall(r"[A-Za-zÅÄÖåäö]", prose))
        if alpha_chars == 0:
            continue
        weighted_scores[lang] += alpha_chars
    if not any(weighted_scores.values()):
        return "finnish"
    return max(weighted_scores, key=weighted_scores.get)

def infer_chat_emotion(line: str, emojis: list[str]) -> str:
    combo = "".join(emojis); lowered = line.lower()
    if any(x in lowered for x in ["wtf", "what the fuck", "holy shit", "jesus", "can't be real"]):
        return "shocked"
    if any(x in lowered for x in ["clip", "content", "let's go", "drop it", "🔥"]):
        return "excited"
    if any(x in lowered for x in ["lol", "lmao", "bro", "😂", "💀"]):
        return "mocking"
    if any(x in lowered for x in ["end the stream", "call police", "stop", "don't"]):
        return "urgent"
    if "?" in line:
        return "questioning"
    if any(x in combo for x in ["😂", "😏", "🙃", "😼"]) or "lol" in lowered: return "mocking"
    if any(x in combo for x in ["😨", "😱", "😰", "😬"]) or "apua" in lowered: return "fearful"
    if any(x in combo for x in ["😡", "🤬", "👿"]): return "angry"
    if any(x in combo for x in ["😍", "❤️", "🥰"]): return "excited"
    return "neutral"

def infer_pace_and_tone(emotion: str, is_chat: bool) -> tuple[str, str]:
    emotion_key = emotion.strip().lower(); pace,tone="medium","neutral"
    if emotion_key in {"angry","viha","vihainen"}: pace,tone="fast","intense"
    elif emotion_key in {"fearful","pelokas"}: pace,tone="medium","tense"
    elif emotion_key in {"excited","innostunut","riemukas"}: pace,tone="fast","bright"
    elif emotion_key in {"mocking","ironinen"}: pace,tone="medium","playful"
    elif emotion_key in {"shocked"}: pace,tone="fast","sharp"
    elif emotion_key in {"urgent"}: pace,tone="fast","alarmed"
    elif emotion_key in {"questioning"}: pace,tone="fast","questioning"
    elif emotion_key in {"surullinen","sad"}: pace,tone="slow","soft"
    if is_chat and pace=="medium" and tone=="neutral": pace,tone="fast","light"
    return pace,tone


def infer_fragment_delivery(text: str, speaker: str, language_name: str, is_chat: bool) -> tuple[str, str, str]:
    """Return emotion, pace and tone for a spoken fragment."""
    lowered = (text or "").lower()
    if language_name == "finnish":
        if is_chat:
            if any(x in lowered for x in ["wtf", "mitä helvettiä", "natsikamaa", "sairasta", "ei vittu", "jumalauta"]) or "?!" in lowered or lowered.count("!") >= 2:
                return "järkyttynyt", "fast", "sharp"
            if any(x in lowered for x in ["gg", "wp", "let's go", "tulessa", "kova", "äijä", "ei ole ihan paskempi", "🔥"]):
                return "innostunut", "fast", "excited"
            if any(x in lowered for x in ["lol", "lmao", "ihan varmasti", "jep jep", "cope", "skill issue"]):
                return "ivallinen", "fast", "mocking"
            if "?" in lowered:
                return "epäuskoinen", "fast", "questioning"
            return "reaktiivinen", "fast", "chatty"
        if speaker == NARRATOR_NAME:
            if any(x in lowered for x in ["jännite", "ase", "veri", "huusi", "räjähdys", "pakeni"]):
                return "jännitteinen", "medium", "tense"
            return "hillitty", "medium", "dry"
        if any(x in lowered for x in ["teidän on nyt syytä", "heti", "tulkaa tänne"]):
            return "päättäväinen", "medium", "firm"
        if any(x in lowered for x in ["pelk", "kauhu", "järky", "apua", "ei voi olla"]):
            return "järkyttynyt", "medium", "tense"
        pace = "fast" if "!" in lowered else "medium"
        tone = "questioning" if "?" in lowered else "neutral"
        return "neutraali", pace, tone
    emotion = infer_chat_emotion(text, []) if is_chat else "neutral"
    pace, tone = infer_pace_and_tone(emotion, is_chat)
    return emotion, pace, tone


def build_delivery_instructions(base_instructions: str, text: str, speaker: str, language_name: str, emotion: str, pace: str, tone: str, is_chat: bool) -> str:
    suffix = ""
    if language_name == "finnish" and is_chat:
        if emotion == "innostunut":
            suffix = "Tämä kommentti on innostunut Twitch-chatin reaktio. Lue se nopeasti, energisesti ja hieman voitonriemuisesti."
        elif emotion == "järkyttynyt":
            suffix = "Tämä kommentti on järkyttynyt ja epäuskoinen chat-reaktio. Lue se nopeasti, terävästi ja reaktiivisesti."
        elif emotion == "ivallinen":
            suffix = "Tämä kommentti on ivallinen chat-heitto. Lue se nopeasti, kuivasti ja pilke äänessä."
        elif emotion == "epäuskoinen":
            suffix = "Tämä kommentti on epäuskoinen chat-kysymys. Lue se nopeasti ja kysyvän terävästi."
        else:
            suffix = "Tämä on nopea chat-reaktio. Pidä delivery napakkana ja elävänä."
    elif language_name == "finnish":
        suffix = f"Pidä tämän fragmentin sävy {emotion} ja {tone}, rytmillä {pace}."
    elif language_name == "english":
        if is_chat:
            suffix = f"This is a reactive chat comment. Read it briskly with a {tone} tone and {pace} pace."
        else:
            suffix = f"Keep this fragment {emotion} in mood, {tone} in tone, with a {pace} pace."
    elif language_name == "spanish":
        if is_chat:
            suffix = f"Este comentario es una reacción de chat. Léelo con ritmo {pace} y tono {tone}."
        else:
            suffix = f"Mantén este fragmento con emoción {emotion}, tono {tone} y ritmo {pace}."
    return f"{base_instructions} {suffix}".strip()


def normalise_finnish_chat_pronunciation(text: str) -> str:
    substitutions = {
        "WTF": "wee-tee-äf",
        "GG": "gee-gee",
        "WP": "vee-pee",
        "LOL": "lol",
        "LMAO": "äl-em-aa-oo",
        "OMG": "oo-em-gee",
        "FPS": "äf-pee-äs",
        "RPG": "är-pee-gee",
        "NPC": "än-pee-see",
        "AI": "aa-ii",
        "APS": "aa-pee-äs",
    }
    out = text
    for key, repl in substitutions.items():
        out = re.sub(rf"\b{re.escape(key)}\b", repl, out, flags=re.IGNORECASE)
    return out


def number_to_finnish_0_99(n: int, *, minute: bool = False) -> str:
    ones = ["nolla", "yksi", "kaksi", "kolme", "neljä", "viisi", "kuusi", "seitsemän", "kahdeksan", "yhdeksän"]
    if minute and n == 0:
        return "tasan"
    if minute and 0 < n < 10:
        return f"nolla {ones[n]}"
    if n < 10:
        return ones[n]
    if n == 10:
        return "kymmenen"
    if 11 <= n <= 19:
        return f"{ones[n-10]}toista"
    tens = n // 10
    rem = n % 10
    tens_word = ["", "", "kaksikymmentä", "kolmekymmentä", "neljäkymmentä", "viisikymmentä", "kuusikymmentä", "seitsemänkymmentä", "kahdeksankymmentä", "yhdeksänkymmentä"][tens]
    return tens_word if rem == 0 else f"{tens_word}{ones[rem]}"


def normalise_finnish_numbers_for_tts(text: str) -> str:
    def repl(m: re.Match[str]) -> str:
        lead = m.group(1)
        hour = int(m.group(2))
        minute = int(m.group(3))
        return f"{lead} {number_to_finnish_0_99(hour)} {number_to_finnish_0_99(minute, minute=True)}"
    return re.sub(r"\b([Kk]ello|[Kk]lo)\s+(\d{1,2})\.(\d{2})\b", lambda m: repl(m).replace("klo", "kello"), text)

def dominant_language_from_text(text: str) -> str:
    prose = strip_non_prose_for_language_detection(text)
    return detect_language_name(prose) if prose else "finnish"


def spoken_chapter_number(content: str, language: str) -> str:
    maps = {
        "finnish": {"0": "nolla", "1": "yksi", "2": "kaksi", "3": "kolme", "4": "neljä", "5": "viisi", "6": "kuusi", "7": "seitsemän", "8": "kahdeksan", "9": "yhdeksän", ".": "piste"},
        "english": {"0": "zero", "1": "one", "2": "two", "3": "three", "4": "four", "5": "five", "6": "six", "7": "seven", "8": "eight", "9": "nine", ".": "point"},
        "spanish": {"0": "cero", "1": "uno", "2": "dos", "3": "tres", "4": "cuatro", "5": "cinco", "6": "seis", "7": "siete", "8": "ocho", "9": "nueve", ".": "punto"},
    }
    table = maps.get(language, maps["english"])
    words = [table[ch] for ch in content.strip() if ch in table]
    return " ".join(words) if words else content


def build_chapter_heading(content: str | None, title: str, language: str) -> str:
    number = (content or "").strip()
    if re.fullmatch(r"\d+\.\d+", number):
        spoken = spoken_chapter_number(number, language)
        if language == "finnish":
            return f"Luku {spoken}, {title}."
        if language == "spanish":
            return f"Capítulo {spoken}, {title}."
        return f"Chapter {spoken}, {title}."
    if language == "finnish":
        return f"Luku {title}."
    if language == "spanish":
        return f"Capítulo {title}."
    return f"Chapter {title}."


def _split_overlong_unit(unit: str, max_len: int) -> list[str]:
    if len(unit) <= max_len:
        return [unit]
    pieces: list[str] = []
    for para in [p.strip() for p in re.split(r"\n\s*\n", unit) if p.strip()]:
        if len(para) <= max_len:
            pieces.append(para)
            continue
        sentences = [x.strip() for x in re.split(r"(?<=[.!?…])\s+", para) if x.strip()]
        cur = ""
        for sent in sentences:
            candidate = f"{cur} {sent}".strip() if cur else sent
            if len(candidate) <= max_len:
                cur = candidate
                continue
            if cur:
                pieces.append(cur)
                cur = ""
            if len(sent) <= max_len:
                cur = sent
                continue
            print("Warning: had to force-split an overlong single narrator unit.", file=sys.stderr)
            rem = sent
            while rem:
                if len(rem) <= max_len:
                    cur = rem
                    rem = ""
                else:
                    cut = rem.rfind(" ", 0, max_len)
                    cut = cut if cut > 0 else max_len
                    pieces.append(rem[:cut].strip())
                    rem = rem[cut:].strip()
        if cur:
            pieces.append(cur)
    return [p for p in pieces if p]


def split_single_voice_fragments(text: str, max_len: int = SINGLE_VOICE_MAX_FRAGMENT_LEN) -> list[str]:
    units = [u.strip() for u in re.split(r"\n\s*\n", text.strip()) if u.strip()]
    prepared: list[str] = []
    for u in units:
        prepared.extend(_split_overlong_unit(u, max_len))
    chunks: list[str] = []
    current = ""
    for unit in prepared:
        candidate = f"{current}\n\n{unit}".strip() if current else unit
        if len(candidate) <= max_len:
            current = candidate
        else:
            if current:
                chunks.append(current)
            current = unit
    if current:
        chunks.append(current)
    return chunks


def build_single_voice_segments(section: str, content: str | None, title: str) -> list[ScriptSegment]:
    normalized = section.strip() + "\n"
    language = dominant_language_from_text(normalized)
    heading = build_chapter_heading(content, title, language)
    fragments = split_single_voice_fragments(normalized, max_len=SINGLE_VOICE_MAX_FRAGMENT_LEN)
    segments = [ScriptSegment(segment_id=1, type="narration", speaker=NARRATOR_NAME, text=heading)]
    for i, frag in enumerate(fragments, start=2):
        segments.append(ScriptSegment(segment_id=i, type="narration", speaker=NARRATOR_NAME, text=frag))
    return segments


def inject_chat_segments(base_segments: list[ScriptSegment]) -> list[ScriptSegment]:
    out: list[ScriptSegment] = []
    in_chat_block = False
    for seg in base_segments:
        if seg.type not in {"narration", "dialogue"}:
            out.append(seg)
            continue
        lines = [ln for ln in seg.text.splitlines() if ln.strip()]
        if not lines:
            continue
        text_buffer: list[str] = []
        for line in lines:
            chat = parse_chat_line(line)
            if chat:
                if text_buffer:
                    if in_chat_block:
                        out.append(ScriptSegment(segment_id=0, type="audio_cue", speaker=NARRATOR_NAME, text="notification_end"))
                        out.append(ScriptSegment(segment_id=0, type="narration", speaker=NARRATOR_NAME, text="", confidence=1.0))
                        in_chat_block = False
                    out.append(ScriptSegment(segment_id=0, type=seg.type, speaker=seg.speaker, text="\n".join(text_buffer)))
                    text_buffer = []
                if not in_chat_block:
                    out.append(ScriptSegment(segment_id=0, type="audio_cue", speaker=NARRATOR_NAME, text="notification_lead"))
                    out.append(ScriptSegment(segment_id=0, type="narration", speaker=NARRATOR_NAME, text="", confidence=1.0))
                    in_chat_block = True
                out.append(ScriptSegment(segment_id=0, type="chat_nickname", speaker="chat", text=chat.nickname))
                out.append(ScriptSegment(segment_id=0, type="audio_cue", speaker=NARRATOR_NAME, text="notification"))
                out.append(ScriptSegment(segment_id=0, type="chat_message", speaker="chat", text=chat.message))
            else:
                if in_chat_block:
                    out.append(ScriptSegment(segment_id=0, type="audio_cue", speaker=NARRATOR_NAME, text="notification_end"))
                    out.append(ScriptSegment(segment_id=0, type="narration", speaker=NARRATOR_NAME, text="", confidence=1.0))
                    in_chat_block = False
                text_buffer.append(line)
        if text_buffer:
            out.append(ScriptSegment(segment_id=0, type=seg.type, speaker=seg.speaker, text="\n".join(text_buffer)))
    if in_chat_block:
        out.append(ScriptSegment(segment_id=0, type="audio_cue", speaker=NARRATOR_NAME, text="notification_end"))
    final: list[ScriptSegment] = []
    for i, seg in enumerate(out, start=1):
        final.append(ScriptSegment(segment_id=i, type=seg.type, speaker=seg.speaker, text=seg.text, confidence=seg.confidence))
    return final


def split_plain_text(text: str, limit: int = MAX_FRAGMENT_LEN) -> list[str]:
    text = text.strip()
    if len(text) <= limit: return [text] if text else []
    parts=[]; rem=text
    while rem:
        if len(rem)<=limit: parts.append(rem.strip()); break
        cut=rem.rfind(" ",0,limit); cut=cut if cut>0 else limit
        parts.append(rem[:cut].strip()); rem=rem[cut:].strip()
    return [p for p in parts if p]

def _longest_alias_prefix(line: str, alias_map: dict[str, str]) -> str | None:
    stripped = line.strip()
    best: str | None = None
    best_len = 0
    for alias_key, canonical in alias_map.items():
        alias_words = [w for w in re.split(r"[^A-Za-zÅÄÖåäö0-9]+", alias_key) if w]
        if not alias_words:
            continue
        alias_text = " ".join(alias_words)
        if stripped.lower().startswith(alias_text.lower()) and len(alias_text) > best_len:
            best = canonical
            best_len = len(alias_text)
    return best


def attribute_speakers_rule_based(text: str, narrators: dict[str, str], alias_map: dict[str, str], debug: bool = False) -> list[ScriptSegment]:
    parsed: list[ScriptSegment] = []
    state={"current_scene_subject":NARRATOR_NAME,"last_explicit_speaker":NARRATOR_NAME,"last_dialogue_speaker":NARRATOR_NAME,"pending_speaker_from_lead_in":NARRATOR_NAME}
    speech_pat = "|".join(sorted(SPEECH_VERBS, key=len, reverse=True))
    for idx, raw in enumerate(text.splitlines(), start=1):
        line=raw.strip()
        if not line: continue
        speaker=NARRATOR_NAME; spoken=line; seg_type="narration"
        m = re.match(r"^([^:]{1,64}):\s*(.*)$", line)
        if m:
            cand=m.group(1).strip(); resolved=resolve_voice_name(cand,narrators,alias_map); seg_type="dialogue"
            if resolved!=NARRATOR_NAME: speaker,spoken=resolved,m.group(2).strip()
        elif line.startswith("–") or line.startswith("“"):
            seg_type="dialogue"; spoken=line.lstrip("–“").strip(); speaker=state["pending_speaker_from_lead_in"]
            end = re.search(rf",\s*([A-ZÅÄÖa-zåäö][^,.:;!?]{{0,40}})\s+({speech_pat})\b", spoken, re.IGNORECASE)
            if end:
                tag_speaker = resolve_voice_name(end.group(1), narrators, alias_map)
                speaker = tag_speaker if tag_speaker != NARRATOR_NAME else speaker
                spoken = re.sub(rf",\s*[A-ZÅÄÖa-zåäö][^,.:;!?]{{0,40}}\s+({speech_pat})\b.*$", "", spoken, flags=re.IGNORECASE).strip()
            elif re.search(rf"\bhän\s+({speech_pat})\b", spoken, re.IGNORECASE) and state["last_explicit_speaker"] != NARRATOR_NAME:
                speaker = state["last_explicit_speaker"]
            if speaker == NARRATOR_NAME:
                speaker = state["last_dialogue_speaker"] if state["last_dialogue_speaker"] != NARRATOR_NAME else NARRATOR_NAME
        else:
            lead = re.search(rf"^([A-ZÅÄÖa-zåäö][^:]]{{0,160}})\b({speech_pat})\b.*:\s*$", line, re.IGNORECASE)
            if lead:
                prefix_match = _longest_alias_prefix(line, alias_map)
                if prefix_match:
                    state["pending_speaker_from_lead_in"] = prefix_match
                else:
                    first_token = re.match(r"^([A-ZÅÄÖa-zåäö][A-Za-zÅÄÖåäö-]{1,40})", line)
                    state["pending_speaker_from_lead_in"] = resolve_voice_name(first_token.group(1) if first_token else lead.group(1), narrators, alias_map)
            else:
                state["pending_speaker_from_lead_in"] = NARRATOR_NAME
        if speaker != NARRATOR_NAME:
            state["last_explicit_speaker"] = speaker; state["last_dialogue_speaker"] = speaker; state["current_scene_subject"] = speaker
        segment = ScriptSegment(segment_id=idx, type=seg_type, speaker=speaker, text=spoken, confidence=1.0)
        parsed.append(segment)
    return parsed


def attribute_speakers_with_openai(text: str, narrators: dict[str, str], alias_map: dict[str, str], model: str = "gpt-4.1-mini", debug: bool = False) -> list[ScriptSegment]:
    client = OpenAI()
    segments = [ln.strip() for ln in text.splitlines() if ln.strip()]
    payload = [{"segment_id": i + 1, "text": s} for i, s in enumerate(segments)]
    narrator_names = ", ".join(narrators.keys())
    prompt = (
        "This is a Finnish novel script. En dash dialogue lines are dialogue. "
        "Dialogue tags such as 'hän sanoi', 'Virtanen vastasi', and 'vaimo huikkasi keittiöstä' are not spoken text. "
        "Infer speaker from the same line and nearby context. "
        f"Use only these narrator names or aliases: {narrator_names}. "
        f"If unclear, use {NARRATOR_NAME}. Return JSON array with fields segment_id,type,speaker,text,confidence."
    )
    res = client.responses.create(model=model, input=[{"role": "system", "content": prompt}, {"role": "user", "content": json.dumps(payload, ensure_ascii=False)}])
    text_out = res.output_text.strip()
    data = json.loads(text_out)
    out: list[ScriptSegment] = []
    for row in data:
        raw_speaker = str(row.get("speaker", NARRATOR_NAME))
        seg = ScriptSegment(
            segment_id=int(row["segment_id"]),
            type="dialogue" if row.get("type") == "dialogue" else "narration",
            speaker=raw_speaker,
            text=str(row.get("text", "")),
            confidence=float(row.get("confidence", 1.0)),
        )
        out.append(seg)
    return out

def to_ssml_lines(
    segments: list[ScriptSegment],
    cfg: NarratorConfig,
    openai_profile_variant: str | None = None,
) -> list[str]:
    out=[]
    default_language = dominant_language_from_segments(segments)
    print(f"[LANG] dominant default_language={default_language}")
    alphabetic_fragment_count = 0
    opposite_language_count = 0
    suspicious_fragments: list[str] = []
    for i,row in enumerate(segments):
        if row.type == "audio_cue":
            out.append(f'<audio src="{html.escape(row.text, quote=True)}"/>')
            continue
        if row.type == "chat_nickname":
            speaker = "chat"
            spoken = EMOJI_RE.sub("", row.text).strip()
            out.append(f'<voice name="{html.escape(speaker, quote=True)}" language="{html.escape(default_language, quote=True)}" emotion="neutral" pace="fast" tone="nickname">{html.escape(spoken)}</voice>')
            continue
        speaker=resolve_voice_name(row.speaker, cfg.voices, cfg.aliases)
        emojis=EMOJI_RE.findall(row.text); spoken=EMOJI_RE.sub("",row.text).strip()
        is_chat=speaker.startswith("chat") or row.type == "chat_message"
        emotion, pace, tone = infer_fragment_delivery(spoken, speaker, default_language, is_chat)
        for chunk_idx, chunk in enumerate(split_plain_text(spoken), start=1):
            cleaned_chunk = strip_non_prose_for_language_detection(chunk)
            alpha_chars = len(re.findall(r"[A-Za-zÅÄÖåäö]", cleaned_chunk))
            detected_language = detect_language_name(cleaned_chunk) if cleaned_chunk else default_language
            if alpha_chars < 80:
                final_language = default_language
                reason = f"short_fragment_alpha<{80}: inherited default_language"
            elif detected_language != default_language and alpha_chars < 160:
                final_language = default_language
                reason = "weak_cross_language_signal: inherited default_language"
            else:
                final_language = detected_language
                reason = "strong_evidence_from_fragment_text"

            emotion, pace, tone = infer_fragment_delivery(chunk, speaker, final_language, is_chat)
            tts_chunk = chunk
            if final_language == "finnish" and is_chat:
                tts_chunk = normalise_finnish_chat_pronunciation(tts_chunk)
            if final_language == "finnish":
                tts_chunk = normalise_finnish_numbers_for_tts(tts_chunk)
            profile = (
                resolve_openai_profile(speaker, final_language, cfg, variant=openai_profile_variant)
                or resolve_openai_profile(speaker, default_language, cfg, variant=openai_profile_variant)
            )
            attrs = [
                f'name="{html.escape(speaker, quote=True)}"',
                f'language="{html.escape(final_language, quote=True)}"',
                f'emotion="{html.escape(emotion, quote=True)}"',
                f'pace="{html.escape(pace, quote=True)}"',
                f'tone="{html.escape(tone, quote=True)}"',
            ]
            profile_name = profile.profile_id if profile else "-"
            preview = re.sub(r"\s+", " ", chunk).strip()[:90]
            print(
                f"[LANG DEBUG] fragment={i+1}.{chunk_idx} speaker={speaker} text={preview!r} "
                f"detected_language={detected_language} inherited/default_language={default_language} "
                f"final_language={final_language} selected_profile={profile_name} reason={reason}"
            )
            if alpha_chars > 0:
                alphabetic_fragment_count += 1
                if (default_language == "english" and final_language == "finnish") or (default_language == "finnish" and final_language == "english"):
                    opposite_language_count += 1
                    suspicious_fragments.append(f"{i+1}.{chunk_idx} speaker={speaker} lang={final_language} text={preview!r}")
            if profile:
                instructions = build_delivery_instructions(profile.instructions, tts_chunk, speaker, final_language, emotion, pace, tone, is_chat)
                attrs.extend(
                    [
                        f'openai_profile="{html.escape(profile.profile_id, quote=True)}"',
                        f'openai_voice="{html.escape(profile.voice, quote=True)}"',
                        f'openai_model="{html.escape(profile.model, quote=True)}"',
                        f'openai_instructions="{html.escape(instructions, quote=True)}"',
                    ]
                )
            out.append(f'<voice {" ".join(attrs)}>{html.escape(tts_chunk)}</voice>')
        if row.type == "chat_message":
            out.append('<break time="0.20s"/>')
        elif i < len(segments)-1:
            out.append(f'<break time="{INTER_LINE_BREAK}"/>')
    if alphabetic_fragment_count > 0 and default_language in {"english", "finnish"}:
        ratio = opposite_language_count / alphabetic_fragment_count
        if ratio > 0.10:
            print(
                f"WARNING: Dominant language is {default_language} but {opposite_language_count}/{alphabetic_fragment_count} "
                f"({ratio:.1%}) alphabetic fragments were marked as the opposite language.",
                file=sys.stderr,
            )
            for frag in suspicious_fragments:
                print(f"  suspicious: {frag}", file=sys.stderr)
    return out

def final_voice_gate(ssml: str, narrators: dict[str, str]) -> str:
    def repl(m: re.Match[str]) -> str:
        val = m.group(1)
        if val in narrators: return m.group(0)
        print(f"Warning: invalid voice name replaced with Kertoja: {val}", file=sys.stderr)
        return m.group(0).replace(val, NARRATOR_NAME)
    return re.sub(r'<voice\s+name="([^"]*)"', repl, ssml)


def normalise_content_for_filename(content: str | None) -> str:
    if not content:
        return "output"
    value = content.strip()
    if re.fullmatch(r"\d+\.\d+", value):
        major, minor = value.split(".", maxsplit=1)
        return f"{int(major):02d}_{int(minor):02d}"
    cleaned = re.sub(r"[^0-9A-Za-z_-]+", "_", value).strip("_")
    return cleaned or "output"


def slugify_filename_part(value: str | None) -> str:
    cleaned = re.sub(r"[^0-9A-Za-zÅÄÖåäö_-]+", "_", (value or "").strip()).strip("_").lower()
    return cleaned or "audiobook"


def resolve_output_path(output_arg: str | None, content: str | None, input_path: Path, chapter_title: str | None = None) -> Path:
    name_root = normalise_content_for_filename(content)
    title_slug = slugify_filename_part(chapter_title) if chapter_title else "audiobook"
    auto_file_name = f"{name_root}_{title_slug}_ssml.xml"
    if not output_arg:
        return Path(auto_file_name)

    output_path = Path(output_arg)
    output_looks_like_dir = output_arg.endswith(("/", "\\"))
    if output_path.exists() and output_path.is_dir():
        output_looks_like_dir = True
    if not output_path.exists() and output_path.suffix == "":
        output_looks_like_dir = True

    if output_looks_like_dir:
        output_path.mkdir(parents=True, exist_ok=True)
        return output_path / auto_file_name

    output_path.parent.mkdir(parents=True, exist_ok=True)
    return output_path


def resolve_code_path(path_arg: str, code_directory: Path) -> Path:
    p = Path(path_arg).expanduser()
    return p if p.is_absolute() else (code_directory / p)


def resolve_work_path(path_arg: str, work_directory: Path) -> Path:
    p = Path(path_arg).expanduser()
    return p if p.is_absolute() else (work_directory / p)


class HelpFormatter(
    argparse.ArgumentDefaultsHelpFormatter,
    argparse.RawDescriptionHelpFormatter,
):
    pass


def main() -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Extract one chapter from a DOCX manuscript and generate SSML for audiobook rendering. "
            "By default the script creates long single-narrator fragments. "
            "Speaker-specific/polyphonic SSML is available with --use-multipolyfony true."
        ),
        formatter_class=HelpFormatter,
        epilog="""Examples:
  Single-narrator SSML, normal production mode:
    python3 tee_aanikasikirjoitus.py \\
      --input manuscript.docx \\
      --content 2.5 \\
      --output /tmp/abook \\
      --narrators-file prompt_narrators.txt

  Polyphonic / speaker-specific SSML:
    python3 tee_aanikasikirjoitus.py \\
      --input manuscript.docx \\
      --content 2.5 \\
      --output /tmp/abook \\
      --narrators-file prompt_narrators.txt \\
      --use-multipolyfony true \\
      --speaker-detection openai \\
      --debug-speakers""",
    )
    required = parser.add_argument_group("required input")
    required.add_argument("--input", required=True, metavar="DOCX", help="Input manuscript .docx file. Relative paths are resolved against --work-directory.")

    chapter = parser.add_argument_group("chapter selection")
    chapter.add_argument("--content", metavar="ACT.CHAPTER", default=None, help="Chapter identifier to extract, for example 2.5. Optional, but normally used for real manuscript chapters. If omitted, the whole document is used.")

    paths = parser.add_argument_group("paths")
    paths.add_argument("--output", metavar="PATH", default=None, help="Output SSML file or output directory. If a directory is given, the file name is generated from --content and the chapter title. Relative paths are resolved against --work-directory.")
    paths.add_argument("--code-directory", metavar="DIR", default=".", help="Directory for support files such as prompt_narrators.txt.")
    paths.add_argument("--work-directory", metavar="DIR", default=".", help="Base directory for manuscript input and SSML output paths.")
    paths.add_argument("--narrators-file", metavar="JSON", default="prompt_narrators.txt", help="Narrator/voice configuration JSON. Relative paths are resolved against --code-directory.")

    mode = parser.add_argument_group("generation mode")
    mode.add_argument("--use-multipolyfony", type=parse_bool_arg, default=False, metavar="BOOL", help="Use speaker-specific/polyphonic SSML. false creates long single-narrator Kertoja fragments and is the normal production mode. Accepted values: true/false, yes/no, 1/0, on/off.")

    speaker = parser.add_argument_group("speaker detection, used mainly with --use-multipolyfony true")
    speaker.add_argument("--speaker-detection", choices=["openai", "rule_based"], default="openai", help="Speaker attribution method for polyphonic SSML.")
    speaker.add_argument("--debug-speakers", action="store_true", help="Print detected segments, speakers, language and OpenAI profile information.")
    speaker.add_argument("--openai-model", default="gpt-4.1-mini", help="OpenAI model used for speaker attribution.")
    speaker.add_argument("--openai-profile-variant", default=None, metavar="VARIANT", help="Optional narrator profile variant, for example fi_ash or fi_fable, if such variants exist in prompt_narrators.txt.")
    args = parser.parse_args()
    try:
        code_directory = Path(args.code_directory).expanduser()
        work_directory = Path(args.work_directory).expanduser()
        narrators_path = resolve_code_path(args.narrators_file, code_directory)
        input_path = resolve_work_path(args.input, work_directory)
        cfg=load_narrator_config(narrators_path)
        section,act_no,chapter_no,title=extract_section(input_path, args.content)
    except Exception as e:
        print(f"Virhe: {e}", file=sys.stderr); return 2
    print(f"Detected chapter title: {title}")
    print(f"Speaker detection mode: {args.speaker_detection}")
    print(f"Use multipolyfony: {args.use_multipolyfony}")
    if args.use_multipolyfony:
        if args.speaker_detection == "openai":
            try:
                segments = attribute_speakers_with_openai(section, cfg.voices, cfg.aliases, model=args.openai_model, debug=args.debug_speakers)
            except Exception as e:
                print(f"WARNING: OpenAI speaker attribution failed: {e}. Falling back to rule-based speaker detection.", file=sys.stderr)
                segments = attribute_speakers_rule_based(section, cfg.voices, cfg.aliases, debug=args.debug_speakers)
        else:
            segments = attribute_speakers_rule_based(section, cfg.voices, cfg.aliases, debug=args.debug_speakers)
        segments = normalise_segments(segments, cfg.voices, cfg.aliases)
        segments = inject_chat_segments(segments)
    else:
        segments = inject_chat_segments(build_single_voice_segments(section, args.content, title))
    if args.debug_speakers:
        for segment in segments:
            conf = f" confidence={segment.confidence:.2f}" if segment.type == "dialogue" else ""
            language = detect_language_name(segment.text)
            profile = resolve_openai_profile(segment.speaker, language, cfg, variant=args.openai_profile_variant)
            profile_text = profile.profile_id if profile else "-"
            print(f"[{segment.type}] speaker={segment.speaker} profile={profile_text}{conf} text={segment.text!r}")
    ssml = "<speak>\n" + "\n".join(
        to_ssml_lines(segments, cfg, openai_profile_variant=args.openai_profile_variant)
    ) + "\n</speak>\n"
    ssml = final_voice_gate(ssml, cfg.voices)
    output_arg = str(resolve_work_path(args.output, work_directory)) if args.output else None
    out = resolve_output_path(output_arg, args.content, input_path, title)
    print(f"Writing SSML to: {out}")
    out.write_text(ssml, encoding="utf-8")
    print(f"Kirjoitettu: {out}")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
